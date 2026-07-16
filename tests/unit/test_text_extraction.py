"""Unit tests for the context-document text extractor.

Covers the dispatch matrix (extension first, MIME fallback), each supported
format (PDF / DOCX / Markdown / plain text), the soft-fail contract for
unsupported and corrupt inputs, storage-cap truncation, and the heuristic
summary derivation.
"""

import io
import subprocess
import tracemalloc
import zlib
from types import SimpleNamespace

import docx
import pytest
from docx import Document as DocxDocument
from pypdf import PasswordType

import apex.services.text_extraction as extraction_module
from apex.services.text_extraction import (
    PARSE_FAILED,
    PARSE_PARSED,
    PARSE_UNSUPPORTED,
    _extract_docx_in_worker,
    _extract_pdf_in_worker,
    _kind,
    derive_summary,
    extract_text,
)


def _make_pdf_stream(content: bytes, *, compressed: bool = False) -> bytes:
    """Build a minimal single-page PDF around one content stream.

    Offsets in the xref table are computed against the assembled bytes so pypdf
    parses it without falling back to object scanning.
    """
    if compressed:
        content = zlib.compress(content)
        filter_entry = b" /Filter /FlateDecode"
    else:
        filter_entry = b""
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length "
        + str(len(content)).encode()
        + filter_entry
        + b" >>\nstream\n"
        + content
        + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{index} 0 obj\n".encode() + obj + b"\nendobj\n"
    xref_pos = len(out)
    out += b"xref\n"
    out += f"0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for offset in offsets:
        out += f"{offset:010d} 00000 n \n".encode()
    out += b"trailer\n"
    out += f"<< /Size {len(objects) + 1} /Root 1 0 R >>\n".encode()
    out += b"startxref\n"
    out += f"{xref_pos}\n".encode()
    out += b"%%EOF"
    return bytes(out)


def _make_pdf(text: str) -> bytes:
    content = f"BT /F1 24 Tf 72 700 Td ({text}) Tj ET".encode("latin-1")
    return _make_pdf_stream(content)


def _make_docx(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    document = DocxDocument()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for c, value in enumerate(row):
                table.cell(r, c).text = value
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.mark.parametrize(
    ("filename", "content_type", "expected"),
    [
        ("spec.pdf", None, "pdf"),
        ("spec.PDF", None, "pdf"),
        ("notes.docx", None, "docx"),
        ("readme.md", None, "text"),
        ("readme.markdown", None, "text"),
        ("plain.txt", None, "text"),
        # Extension wins over a mismatched content type.
        ("doc.pdf", "text/plain", "pdf"),
        # No useful extension -> fall back to MIME.
        ("blob", "application/pdf", "pdf"),
        (
            "blob",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "docx",
        ),
        ("blob", "text/markdown", "text"),
        ("blob", "text/csv", "text"),  # any text/* is treated as plain text
        ("archive.zip", "application/zip", None),
        ("image.png", "image/png", None),
    ],
)
def test_kind_dispatch(filename: str, content_type: str | None, expected: str | None) -> None:
    assert _kind(filename, content_type) == expected


def test_extract_plaintext_normalizes_newlines() -> None:
    data = b"first line\r\nsecond line  \r\n\r\n"
    result = extract_text(data, filename="a.txt", content_type=None, max_chars=10_000)
    assert result.status == PARSE_PARSED
    assert result.text == "first line\nsecond line"
    assert result.char_count == len("first line\nsecond line")


def test_extract_markdown_passthrough() -> None:
    data = b"# Title\n\nSome **bold** evidence.\n"
    result = extract_text(data, filename="story.md", content_type=None, max_chars=10_000)
    assert result.status == PARSE_PARSED
    assert "# Title" in result.text
    assert "**bold**" in result.text


def test_extract_plaintext_lossy_fallback_on_invalid_utf8() -> None:
    data = b"valid \xff\xfe bytes"
    result = extract_text(data, filename="weird.txt", content_type=None, max_chars=10_000)
    assert result.status == PARSE_PARSED
    assert "valid" in result.text
    assert "bytes" in result.text


def test_plaintext_normalization_is_chunk_bounded_for_newline_bomb() -> None:
    # The old whole-document split created one Python list entry per newline
    # (hundreds of MiB at the 25 MiB upload limit). Input is allocated before
    # tracing so the assertion measures extraction overhead only.
    data = b"\n" * (8 * 1024 * 1024)
    tracemalloc.start()
    try:
        result = extract_text(data, filename="newlines.txt", content_type=None, max_chars=200_000)
        _current, peak = tracemalloc.get_traced_memory()
    finally:
        tracemalloc.stop()

    assert result.status == PARSE_PARSED
    assert result.text == ""
    assert result.char_count == 0
    assert peak < 2 * 1024 * 1024


def test_plaintext_normalizes_crlf_split_across_decode_chunks() -> None:
    prefix = b"a" * (extraction_module._DECODE_CHUNK_BYTES - 1)
    result = extract_text(
        prefix + b"\r\nnext  \rfinal",
        filename="boundary.txt",
        content_type=None,
        max_chars=100_000,
    )
    expected = prefix.decode() + "\nnext\nfinal"
    assert result.text == expected
    assert result.char_count == len(expected)


def test_extract_docx_paragraphs_and_tables() -> None:
    data = _make_docx(
        ["Heading paragraph", "Second paragraph"],
        table_rows=[["r1c1", "r1c2"], ["r2c1", "r2c2"]],
    )
    result = extract_text(data, filename="notes.docx", content_type=None, max_chars=10_000)
    assert result.status == PARSE_PARSED
    assert "Heading paragraph" in result.text
    assert "Second paragraph" in result.text
    assert "r1c1\tr1c2" in result.text
    assert "r2c1\tr2c2" in result.text


def test_docx_feeds_every_large_block_without_retaining_block_list(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    # Zip validation still receives a real DOCX, while the document object makes
    # block sizes deterministic and proves later blocks are visited after the
    # retained prefix is full.
    valid_docx = _make_docx(["stub"])
    paragraphs = [SimpleNamespace(text="a" * 200_000), SimpleNamespace(text="tail")]
    fake_document = SimpleNamespace(paragraphs=paragraphs, tables=[])
    monkeypatch.setattr(docx, "Document", lambda _stream: fake_document)

    extracted = _extract_docx_in_worker(valid_docx, max_chars=64)

    assert extracted.text == "a" * 64
    assert extracted.char_count == 200_005


def test_extract_pdf_returns_text() -> None:
    data = _make_pdf("Hello PDF context")
    result = extract_text(data, filename="spec.pdf", content_type=None, max_chars=10_000)
    assert result.status == PARSE_PARSED
    assert "Hello" in result.text
    assert "context" in result.text


def test_pdf_nonraising_failed_decryption_is_rejected(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    class EncryptedReader:
        is_encrypted = True

        def __init__(self, _stream: object) -> None:
            pass

        def decrypt(self, _password: str) -> PasswordType:
            return PasswordType.NOT_DECRYPTED

        @property
        def pages(self) -> object:
            raise AssertionError("encrypted pages must not be inspected")

    monkeypatch.setattr("pypdf.PdfReader", EncryptedReader)

    with pytest.raises(extraction_module._ExtractionPolicyError, match="password-protected"):
        _extract_pdf_in_worker(b"%PDF-encrypted", max_chars=100)


def test_pdf_rejects_small_compressed_input_with_oversized_decoded_stream() -> None:
    decoded = b" " * (extraction_module._MAX_PDF_DECODED_STREAM_BYTES + 1)
    data = _make_pdf_stream(decoded, compressed=True)
    assert len(data) < 100_000

    result = extract_text(data, filename="bomb.pdf", content_type=None, max_chars=10_000)

    assert result.status == PARSE_FAILED
    assert result.text == ""
    assert result.error is not None
    assert "limit" in result.error.lower()


@pytest.mark.parametrize(
    ("filename", "data", "worker_module"),
    [
        ("isolated.pdf", b"%PDF", "apex.services.pdf_extraction_worker"),
        ("isolated.docx", b"PK", "apex.services.docx_extraction_worker"),
    ],
)
def test_complex_format_workers_do_not_inherit_server_secrets(
    monkeypatch: pytest.MonkeyPatch,
    filename: str,
    data: bytes,
    worker_module: str,
) -> None:
    sensitive_names = {
        "APEX_AUTH__DEV_API_KEY",
        "AWS_SECRET_ACCESS_KEY",
        "DATABASE_URI",
        "JIRA_TOKEN",
        "REDIS_URI",
    }
    for name in sensitive_names:
        monkeypatch.setenv(name, "sentinel-secret-must-not-reach-worker")

    captured_environment: dict[str, str] | None = None

    class SuccessfulProcess:
        returncode = 0

        def communicate(
            self,
            input: bytes | None = None,
            timeout: float | None = None,
        ) -> tuple[bytes, bytes]:
            del input, timeout
            return b"S" + (0).to_bytes(8, "big"), b""

    def fake_popen(
        command: list[str],
        *,
        stdin: int,
        stdout: int,
        stderr: int,
        env: dict[str, str],
    ) -> SuccessfulProcess:
        nonlocal captured_environment
        del stdin, stdout, stderr
        assert command[1:4] == ["-I", "-m", worker_module]
        captured_environment = dict(env)
        return SuccessfulProcess()

    monkeypatch.setattr(extraction_module.subprocess, "Popen", fake_popen)

    result = extract_text(data, filename=filename, content_type=None, max_chars=100)

    assert result.status == PARSE_PARSED
    assert captured_environment == extraction_module._EXTRACTION_WORKER_ENV
    assert captured_environment is not None
    assert sensitive_names.isdisjoint(captured_environment)


def test_pdf_timeout_kills_and_reaps_worker(monkeypatch: pytest.MonkeyPatch) -> None:
    calls: list[tuple[str, float | None]] = []

    class HungProcess:
        returncode = -9

        def communicate(
            self,
            input: bytes | None = None,
            timeout: float | None = None,
        ) -> tuple[bytes, bytes]:
            del input
            calls.append(("communicate", timeout))
            if len(calls) == 1:
                raise subprocess.TimeoutExpired(cmd="pdf-worker", timeout=timeout or 0)
            return b"", b""

        def kill(self) -> None:
            calls.append(("kill", None))

    monkeypatch.setattr(
        extraction_module.subprocess,
        "Popen",
        lambda *_args, **_kwargs: HungProcess(),
    )

    result = extract_text(b"%PDF", filename="hung.pdf", content_type=None, max_chars=100)

    assert result.status == PARSE_FAILED
    assert result.error is not None and "second limit" in result.error
    assert calls == [
        ("communicate", extraction_module._PDF_WORKER_WALL_SECONDS),
        ("kill", None),
        ("communicate", extraction_module._PDF_WORKER_REAP_SECONDS),
    ]


def test_docx_timeout_kills_and_reaps_isolated_worker(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    calls: list[tuple[str, float | None]] = []

    class HungProcess:
        returncode = -9

        def communicate(
            self,
            input: bytes | None = None,
            timeout: float | None = None,
        ) -> tuple[bytes, bytes]:
            del input
            calls.append(("communicate", timeout))
            if len(calls) == 1:
                raise subprocess.TimeoutExpired(cmd="docx-worker", timeout=timeout or 0)
            return b"", b""

        def kill(self) -> None:
            calls.append(("kill", None))

    monkeypatch.setattr(
        extraction_module.subprocess,
        "Popen",
        lambda *_args, **_kwargs: HungProcess(),
    )

    result = extract_text(
        _make_docx(["untrusted"]),
        filename="hung.docx",
        content_type=None,
        max_chars=100,
    )

    assert result.status == PARSE_FAILED
    assert result.error is not None and "second limit" in result.error
    assert calls == [
        ("communicate", extraction_module._DOCX_WORKER_WALL_SECONDS),
        ("kill", None),
        ("communicate", extraction_module._DOCX_WORKER_REAP_SECONDS),
    ]


@pytest.mark.parametrize(
    ("extractor", "wall_seconds"),
    [
        (extraction_module._extract_pdf, extraction_module._PDF_WORKER_WALL_SECONDS),
        (extraction_module._extract_docx, extraction_module._DOCX_WORKER_WALL_SECONDS),
    ],
)
def test_complex_worker_timeout_does_not_retain_partial_document_output(
    monkeypatch: pytest.MonkeyPatch,
    extractor: object,
    wall_seconds: float,
) -> None:
    canary = b"bare-document-timeout-canary"
    calls = 0

    class HungProcess:
        returncode = -9

        def communicate(self, **_kwargs: object) -> tuple[bytes, bytes]:
            nonlocal calls
            calls += 1
            if calls == 1:
                raise subprocess.TimeoutExpired(
                    cmd="worker",
                    timeout=wall_seconds,
                    output=canary,
                )
            return b"", b""

        def kill(self) -> None:
            return None

    monkeypatch.setattr(extraction_module.subprocess, "Popen", lambda *_a, **_k: HungProcess())
    run = extractor
    assert callable(run)

    with pytest.raises(ValueError, match="second limit") as excinfo:
        run(b"document", max_chars=100)

    assert excinfo.value.__cause__ is None
    assert excinfo.value.__context__ is None
    assert canary.decode() not in str(excinfo.value)


def test_pdf_worker_response_rejects_output_beyond_retention_budget() -> None:
    payload = b"S" + (101).to_bytes(8, "big") + (b"x" * 101)
    with pytest.raises(ValueError, match="output limit"):
        extraction_module._decode_pdf_worker_response(payload, max_chars=25)


def test_unsupported_type_reports_unsupported() -> None:
    result = extract_text(b"\x00\x01\x02", filename="archive.zip", content_type=None, max_chars=100)
    assert result.status == PARSE_UNSUPPORTED
    assert result.text == ""
    assert result.char_count == 0
    assert result.error is None


def test_corrupt_pdf_soft_fails() -> None:
    # A .pdf extension routes to the PDF extractor, but the bytes are not a PDF.
    result = extract_text(
        b"this is not a pdf", filename="broken.pdf", content_type=None, max_chars=100
    )
    assert result.status == PARSE_FAILED
    assert result.text == ""
    assert result.error is not None


def test_extraction_failure_diagnostic_redacts_credentials(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    canary = "parser-secret-canary-9d4a"

    def fail(*_args: object, **_kwargs: object) -> object:
        raise ValueError(f"password={canary}")

    monkeypatch.setattr(extraction_module, "_extract_plaintext", fail)
    result = extract_text(b"body", filename="broken.txt", content_type=None, max_chars=100)

    assert result.status == PARSE_FAILED
    assert result.error is not None
    assert canary not in result.error
    assert "[REDACTED]" in result.error


@pytest.mark.parametrize(
    ("filename", "data", "label"),
    [("hostile.pdf", b"pdf", "PDF"), ("hostile.docx", b"docx", "DOCX")],
)
def test_complex_worker_arbitrary_parser_error_is_not_exposed(
    monkeypatch: pytest.MonkeyPatch,
    filename: str,
    data: bytes,
    label: str,
) -> None:
    canary = "bare-parser-exception-canary-4f01e85b"

    class FailedProcess:
        returncode = 0

        def communicate(self, **_kwargs: object) -> tuple[bytes, bytes]:
            return b"E" + canary.encode(), b""

    monkeypatch.setattr(extraction_module.subprocess, "Popen", lambda *_a, **_k: FailedProcess())

    result = extract_text(data, filename=filename, content_type=None, max_chars=100)

    assert result.status == PARSE_FAILED
    assert result.error == f"ValueError: {label} extraction failed"
    assert canary not in repr(result)


def test_truncation_applies_storage_cap_and_reports_full_length() -> None:
    body = "abcdefghij" * 10  # 100 chars
    result = extract_text(body.encode(), filename="big.txt", content_type=None, max_chars=25)
    assert result.status == PARSE_PARSED
    assert result.char_count == 100
    assert "…[truncated 75 characters]" in result.text
    # The retained slice never exceeds the cap (before the marker).
    assert result.text.split("\n\n…[truncated")[0] == body[:25].rstrip()


def test_no_truncation_when_under_cap() -> None:
    result = extract_text(b"short", filename="s.txt", content_type=None, max_chars=1_000)
    assert result.text == "short"
    assert "truncated" not in result.text


def test_max_chars_zero_disables_truncation() -> None:
    body = ("x" * 500).encode()
    result = extract_text(body, filename="s.txt", content_type=None, max_chars=0)
    assert result.char_count == 500
    assert "truncated" not in result.text


def test_derive_summary_uses_first_nonempty_paragraph() -> None:
    text = "\n\n  \n\nFirst real   paragraph here.\n\nSecond paragraph."
    assert derive_summary(text, max_chars=280) == "First real paragraph here."


def test_derive_summary_caps_with_ellipsis() -> None:
    summary = derive_summary("word " * 100, max_chars=20)
    assert summary is not None
    assert len(summary) == 20
    assert summary.endswith("…")


def test_derive_summary_none_for_blank() -> None:
    assert derive_summary("\n\n   \n\n", max_chars=280) is None


def test_normalized_builder_handles_unlimited_pending_whitespace_and_terminal_cr() -> None:
    builder = extraction_module._NormalizedTextBuilder(None)
    builder.feed("")
    builder.feed("  first  ")
    builder.feed("\n   second\r")
    # A terminal CR is resolved even when the next feed contains only its LF.
    builder.feed("\n")
    builder.feed("third")

    result = builder.finish()

    assert result.text == "first\n   second\nthird"
    assert result.char_count == len(result.text)


def test_normalized_builder_counts_content_after_retention_is_full() -> None:
    builder = extraction_module._NormalizedTextBuilder(3)
    builder.feed("abc   \nnext")

    result = builder.finish()

    assert result.text == "abc"
    assert result.char_count == len("abc\nnext")


def test_normalized_builder_empty_internal_segments_are_noops() -> None:
    builder = extraction_module._NormalizedTextBuilder(0)
    builder._add_outer_pending("")
    builder._append_committed("")
    builder._confirm_non_whitespace("   ")
    builder._feed_chunk("")
    assert builder.finish() == extraction_module._BoundedText(text="", char_count=0)


def test_configure_pypdf_limits_enforces_stream_and_array_ceilings(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from pypdf import filters
    from pypdf.generic import _data_structures

    monkeypatch.setattr(filters, "FLATE_MAX_BUFFER_SIZE", 0)
    monkeypatch.setattr(filters, "JBIG2_MAX_OUTPUT_LENGTH", 1)
    monkeypatch.setattr(
        _data_structures,
        "CONTENT_STREAM_ARRAY_MAX_LENGTH",
        extraction_module._MAX_PDF_CONTENT_STREAMS + 100,
    )

    extraction_module._configure_pypdf_limits()

    assert filters.FLATE_MAX_BUFFER_SIZE == extraction_module._MAX_PDF_DECODED_STREAM_BYTES
    assert filters.JBIG2_MAX_OUTPUT_LENGTH == 1
    assert (
        _data_structures.CONTENT_STREAM_ARRAY_MAX_LENGTH
        == extraction_module._MAX_PDF_CONTENT_STREAMS
    )


def test_pdf_worker_parser_normalizes_multiple_pages_in_process() -> None:
    extracted = extraction_module._extract_pdf_in_worker(
        _make_pdf("in-process parser"), max_chars=10_000
    )

    assert "in-process parser" in extracted.text
    assert extracted.char_count == len(extracted.text)


def test_pdf_worker_rejects_encrypted_document_when_blank_decrypt_fails(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import pypdf

    class EncryptedReader:
        is_encrypted = True
        pages: list[object] = []

        def __init__(self, _stream: object) -> None:
            pass

        def decrypt(self, _password: str) -> None:
            raise RuntimeError("secret details")

    monkeypatch.setattr(pypdf, "PdfReader", EncryptedReader)

    with pytest.raises(ValueError, match="password-protected"):
        extraction_module._extract_pdf_in_worker(b"pdf", max_chars=100)


def test_pdf_worker_enforces_page_and_decoded_content_budgets(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import pypdf

    class Contents:
        def get_data(self) -> bytes:
            return b"oversized"

    class Page:
        def get_contents(self) -> Contents:
            return Contents()

        def extract_text(self) -> str:
            return "never retained"

    class Reader:
        is_encrypted = False

        def __init__(self, _stream: object) -> None:
            self.pages = [Page(), Page()]

    monkeypatch.setattr(pypdf, "PdfReader", Reader)
    monkeypatch.setattr(extraction_module, "_MAX_PDF_PAGES", 1)
    with pytest.raises(ValueError, match="page extraction limit"):
        extraction_module._extract_pdf_in_worker(b"pdf", max_chars=100)

    monkeypatch.setattr(extraction_module, "_MAX_PDF_PAGES", 2)
    monkeypatch.setattr(extraction_module, "_MAX_PDF_DECODED_CONTENT_BYTES", 1)
    with pytest.raises(ValueError, match="decoded page content"):
        extraction_module._extract_pdf_in_worker(b"pdf", max_chars=100)


def test_pdf_worker_separates_pages_and_stops_after_character_budget(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import pypdf

    class Page:
        def __init__(self, text: str) -> None:
            self._text = text

        def get_contents(self) -> None:
            return None

        def extract_text(self) -> str:
            return self._text

    class Reader:
        is_encrypted = False

        def __init__(self, _stream: object) -> None:
            self.pages = [Page("one"), Page("two"), Page("unvisited")]

    monkeypatch.setattr(pypdf, "PdfReader", Reader)

    extracted = extraction_module._extract_pdf_in_worker(b"pdf", max_chars=5)

    assert extracted.text == "one\n\n"
    assert extracted.char_count == len("one\n\ntwo")


@pytest.mark.parametrize(
    ("decoder", "label"),
    [
        (extraction_module._decode_pdf_worker_response, "PDF"),
        (extraction_module._decode_docx_worker_response, "DOCX"),
    ],
)
@pytest.mark.parametrize(
    ("payload", "message"),
    [
        (b"", "without a result"),
        (b"E", "extraction failed"),
        (b"Xshort", "malformed result"),
        (b"S" + (401).to_bytes(8, "big") + (b"x" * 401), "output limit"),
        (b"S" + (1).to_bytes(8, "big") + b"\xff", "invalid UTF-8"),
        (b"S" + (0).to_bytes(8, "big") + b"x", "inconsistent bounds"),
    ],
)
def test_worker_response_decoders_fail_closed(
    decoder: object,
    label: str,
    payload: bytes,
    message: str,
) -> None:
    decode = decoder
    assert callable(decode)
    with pytest.raises(ValueError, match=f"{label}.*{message}"):
        decode(payload, max_chars=100)


@pytest.mark.parametrize(
    "decoder",
    [
        extraction_module._decode_pdf_worker_response,
        extraction_module._decode_docx_worker_response,
    ],
)
def test_worker_invalid_utf8_does_not_retain_extracted_document_bytes(
    decoder: object,
) -> None:
    canary = b"bare-document-decode-canary"
    payload = b"S" + (len(canary) + 1).to_bytes(8, "big") + b"\xff" + canary
    decode = decoder
    assert callable(decode)

    with pytest.raises(ValueError, match="invalid UTF-8") as excinfo:
        decode(payload, max_chars=100)

    assert excinfo.value.__cause__ is None
    assert excinfo.value.__context__ is None
    assert canary.decode() not in str(excinfo.value)


def test_docx_worker_parser_visits_paragraphs_tables_and_cells_in_process() -> None:
    data = _make_docx(["paragraph"], table_rows=[["left", "right"]])

    extracted = extraction_module._extract_docx_in_worker(data, max_chars=10_000)

    assert extracted.text == "paragraph\nleft\tright"
    assert extracted.char_count == len(extracted.text)


@pytest.mark.parametrize(
    ("member", "message"),
    [
        (SimpleNamespace(file_size=1, compress_size=0), "invalid compressed entry"),
        (SimpleNamespace(file_size=101, compress_size=1), "compression ratio"),
    ],
)
def test_docx_worker_rejects_unsafe_archive_members(
    monkeypatch: pytest.MonkeyPatch,
    member: SimpleNamespace,
    message: str,
) -> None:
    class Archive:
        def __enter__(self) -> "Archive":
            return self

        def __exit__(self, *_args: object) -> None:
            pass

        def infolist(self) -> list[SimpleNamespace]:
            return [member]

    monkeypatch.setattr(extraction_module.zipfile, "ZipFile", lambda _stream: Archive())

    with pytest.raises(ValueError, match=message):
        extraction_module._extract_docx_in_worker(b"docx", max_chars=100)


def test_docx_worker_rejects_entry_count_and_expanded_size(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    data = _make_docx(["content"])
    monkeypatch.setattr(extraction_module, "_MAX_DOCX_ENTRIES", 0)
    with pytest.raises(ValueError, match="too many archive entries"):
        extraction_module._extract_docx_in_worker(data, max_chars=100)

    monkeypatch.setattr(extraction_module, "_MAX_DOCX_ENTRIES", 10_000)
    monkeypatch.setattr(extraction_module, "_MAX_DOCX_EXPANDED_BYTES", 1)
    with pytest.raises(ValueError, match="expanded content"):
        extraction_module._extract_docx_in_worker(data, max_chars=100)


@pytest.mark.parametrize(
    ("extractor", "limit_name", "label"),
    [
        (extraction_module._extract_pdf, "_MAX_PDF_INPUT_BYTES", "PDF"),
        (extraction_module._extract_docx, "_MAX_DOCX_INPUT_BYTES", "DOCX"),
    ],
)
def test_complex_extractors_reject_oversized_input_before_spawning(
    monkeypatch: pytest.MonkeyPatch,
    extractor: object,
    limit_name: str,
    label: str,
) -> None:
    monkeypatch.setattr(extraction_module, limit_name, 1)
    extract = extractor
    assert callable(extract)
    with pytest.raises(ValueError, match=f"{label} exceeds the 1-byte"):
        extract(b"xx", max_chars=100)


@pytest.mark.parametrize(
    ("filename", "data", "label"),
    [("failed.pdf", b"pdf", "PDF"), ("failed.docx", b"docx", "DOCX")],
)
def test_complex_extractor_reports_nonzero_worker_exit(
    monkeypatch: pytest.MonkeyPatch,
    filename: str,
    data: bytes,
    label: str,
) -> None:
    class FailedProcess:
        returncode = 9

        def communicate(self, **_kwargs: object) -> tuple[bytes, bytes]:
            return b"", b""

    monkeypatch.setattr(extraction_module.subprocess, "Popen", lambda *_a, **_k: FailedProcess())

    result = extract_text(data, filename=filename, content_type=None, max_chars=100)

    assert result.status == PARSE_FAILED
    assert result.error is not None and f"{label} extraction exceeded" in result.error


@pytest.mark.parametrize(
    ("filename", "data"),
    [("stuck.pdf", b"pdf"), ("stuck.docx", b"docx")],
)
def test_complex_extractor_retries_reap_after_second_timeout(
    monkeypatch: pytest.MonkeyPatch,
    filename: str,
    data: bytes,
) -> None:
    calls: list[str] = []
    communicate_attempts = 0

    class StuckProcess:
        returncode = -9

        def communicate(self, **_kwargs: object) -> tuple[bytes, bytes]:
            nonlocal communicate_attempts
            communicate_attempts += 1
            calls.append("communicate")
            if communicate_attempts <= 2:
                raise subprocess.TimeoutExpired(cmd="worker", timeout=1)
            return b"", b""

        def kill(self) -> None:
            calls.append("kill")

    monkeypatch.setattr(extraction_module.subprocess, "Popen", lambda *_a, **_k: StuckProcess())

    result = extract_text(data, filename=filename, content_type=None, max_chars=100)

    assert result.status == PARSE_FAILED
    assert calls == ["communicate", "kill", "communicate", "kill", "communicate"]
