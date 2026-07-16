"""Bounded plain-text extraction for uploaded context documents.

PDF and DOCX parsing are isolated in short-lived worker processes because their
third-party parsers consume attacker-controlled compressed/container formats.
Plain-text normalization retains only the configured prefix and never builds a
``splitlines()`` list proportional to the input.
"""

from __future__ import annotations

import codecs
import re
import subprocess
import sys
import zipfile
from dataclasses import dataclass
from io import BytesIO
from posixpath import splitext
from typing import Any

from apex.domain.diagnostics import bounded_diagnostic

# Parse-status values surfaced through the API/UI.
PARSE_PARSED = "parsed"
PARSE_FAILED = "failed"
PARSE_UNSUPPORTED = "unsupported"

# Extensions we know how to turn into text.
_PDF_EXTS = {".pdf"}
_DOCX_EXTS = {".docx"}
_TEXT_EXTS = {".md", ".markdown", ".txt", ".text"}

# MIME fallbacks for when the filename has no useful extension.
_PDF_MIMES = {"application/pdf"}
_DOCX_MIMES = {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
_TEXT_MIMES = {"text/markdown", "text/x-markdown", "text/plain"}

_MAX_PDF_PAGES = 2_000
_MAX_PDF_INPUT_BYTES = 25 * 1024 * 1024
_MAX_PDF_DECODED_STREAM_BYTES = 8 * 1024 * 1024
_MAX_PDF_DECODED_CONTENT_BYTES = 16 * 1024 * 1024
_MAX_PDF_CONTENT_STREAMS = 512
_PDF_WORKER_MEMORY_BYTES = 256 * 1024 * 1024
_PDF_WORKER_CPU_SECONDS = 15
_PDF_WORKER_WALL_SECONDS = 20.0
_PDF_WORKER_REAP_SECONDS = 2.0
_WORKER_ERROR_BYTES = 1_000
_WORKER_ERROR_PARSER = "parser_failure"
_WORKER_ERROR_SAFETY_LIMIT = "safety_limit"
_WORKER_ERROR_PASSWORD_PROTECTED = "password_protected"
_WORKER_ERROR_INPUT_LIMIT = "input_limit"
_WORKER_ERROR_CONFIGURATION = "configuration_failure"

_MAX_DOCX_ENTRIES = 2_048
_MAX_DOCX_INPUT_BYTES = 25 * 1024 * 1024
_MAX_DOCX_EXPANDED_BYTES = 64 * 1024 * 1024
_MAX_DOCX_COMPRESSION_RATIO = 100
_DOCX_WORKER_MEMORY_BYTES = 256 * 1024 * 1024
_DOCX_WORKER_CPU_SECONDS = 15
_DOCX_WORKER_WALL_SECONDS = 20.0
_DOCX_WORKER_REAP_SECONDS = 2.0
_DECODE_CHUNK_BYTES = 64 * 1024
_NORMALIZE_CHUNK_CHARS = 64 * 1024
_UNBOUNDED_COMPLEX_FORMAT_CHARS = 1_000_000
_NEWLINE_RUN_RE = re.compile(r"\n+")

# Parser workers handle attacker-controlled files but do not need any server
# credentials.  Supplying an explicit environment prevents subprocesses from
# inheriting database, provider, signing, and integration secrets.
_EXTRACTION_WORKER_ENV = {
    "LANG": "C",
    "LC_ALL": "C",
}


@dataclass(frozen=True)
class ExtractionResult:
    """Outcome of a single extraction.

    ``text`` is normalized and storage-capped. ``char_count`` is the normalized
    length observed before the storage cap; complex formats stop after their
    hard safety budget rather than continuing unbounded parser work.
    """

    text: str
    status: str
    char_count: int
    error: str | None = None


@dataclass(frozen=True)
class _BoundedText:
    text: str
    char_count: int


class _ExtractionPolicyError(ValueError):
    """Server-owned worker failure with a stable cross-process error code."""

    def __init__(self, worker_code: str, message: str) -> None:
        self.worker_code = worker_code
        super().__init__(message)


class _NormalizedTextBuilder:
    """Normalize incrementally while retaining at most ``max_chars`` characters.

    This exactly implements the previous newline/whitespace contract:

    * CRLF and CR become LF;
    * trailing whitespace is removed from every line; and
    * whitespace around the complete document is removed.

    Pending whitespace is kept separately until a later non-whitespace
    character proves it is internal.  Its retained representation is capped,
    so a multi-megabyte line of spaces/newlines remains constant-memory.
    """

    def __init__(self, max_chars: int | None) -> None:
        self._limit = max_chars if max_chars is not None and max_chars >= 0 else None
        self._parts: list[str] = []
        self._retained_chars = 0
        self._char_count = 0
        self._started = False

        self._outer_pending_count = 0
        self._outer_pending_parts: list[str] = []
        self._outer_pending_retained = 0
        self._line_pending_count = 0
        self._line_pending_parts: list[str] = []
        self._line_pending_retained = 0
        self._pending_cr = False

    @property
    def char_count(self) -> int:
        return self._char_count

    def _retention_room(self) -> int | None:
        if self._limit is None:
            return None
        return max(self._limit - self._retained_chars, 0)

    @staticmethod
    def _prefix_for_pending(text: str, *, room: int | None, retained: int) -> str:
        if room is None:
            return text
        remaining = max(room - retained, 0)
        return text[:remaining]

    def _add_line_pending(self, text: str) -> None:
        if not text:
            return
        self._line_pending_count += len(text)
        prefix = self._prefix_for_pending(
            text,
            room=self._retention_room(),
            retained=self._line_pending_retained,
        )
        if prefix:
            self._line_pending_parts.append(prefix)
            self._line_pending_retained += len(prefix)

    def _add_outer_pending(self, text: str) -> None:
        if not text:
            return
        self._outer_pending_count += len(text)
        prefix = self._prefix_for_pending(
            text,
            room=self._retention_room(),
            retained=self._outer_pending_retained,
        )
        if prefix:
            self._outer_pending_parts.append(prefix)
            self._outer_pending_retained += len(prefix)

    def _discard_line_pending(self) -> None:
        self._line_pending_count = 0
        self._line_pending_parts.clear()
        self._line_pending_retained = 0

    def _discard_outer_pending(self) -> None:
        self._outer_pending_count = 0
        self._outer_pending_parts.clear()
        self._outer_pending_retained = 0

    def _move_line_pending_to_outer(self) -> None:
        if self._line_pending_count == 0:
            return
        self._outer_pending_count += self._line_pending_count
        room = self._retention_room()
        for part in self._line_pending_parts:
            prefix = self._prefix_for_pending(
                part,
                room=room,
                retained=self._outer_pending_retained,
            )
            if prefix:
                self._outer_pending_parts.append(prefix)
                self._outer_pending_retained += len(prefix)
        self._discard_line_pending()

    def _append_committed(self, text: str) -> None:
        if not text:
            return
        self._char_count += len(text)
        room = self._retention_room()
        prefix = text if room is None else text[:room]
        if prefix:
            self._parts.append(prefix)
            self._retained_chars += len(prefix)

    def _commit_outer_pending(self) -> None:
        if self._outer_pending_count == 0:
            return
        self._char_count += self._outer_pending_count
        room = self._retention_room()
        if room is None:
            prefixes = self._outer_pending_parts
        else:
            prefixes: list[str] = []
            remaining = room
            for part in self._outer_pending_parts:
                if remaining <= 0:
                    break
                prefix = part[:remaining]
                prefixes.append(prefix)
                remaining -= len(prefix)
        for prefix in prefixes:
            if prefix:
                self._parts.append(prefix)
                self._retained_chars += len(prefix)
        self._discard_outer_pending()

    def _confirm_non_whitespace(self, text: str) -> None:
        """Commit a segment known to end in a non-whitespace character."""

        self._move_line_pending_to_outer()
        if self._started:
            self._commit_outer_pending()
        else:
            # Overall ``strip()`` removes pending newlines/indentation before the
            # first real character.
            self._discard_outer_pending()
            text = text.lstrip()
            if not text:
                return
            self._started = True
        self._append_committed(text)

    def _feed_line_segment(self, text: str) -> None:
        if not text:
            return
        core = text.rstrip()
        if core:
            self._confirm_non_whitespace(core)
        self._add_line_pending(text[len(core) :])

    def _line_break(self) -> None:
        # Per-line rstrip discards these characters; the normalized LF itself is
        # pending until later content proves it is not overall trailing space.
        self._discard_line_pending()
        self._add_outer_pending("\n")

    def _feed_chunk(self, text: str) -> None:
        if not text:
            return
        if self._pending_cr:
            self._line_break()
            if text.startswith("\n"):
                text = text[1:]
            self._pending_cr = False
            if not text:
                return

        # Defer a terminal CR until the next chunk so a split CRLF remains one
        # newline. Replacements are chunk-bounded, unlike the old whole-document
        # normalization.
        if text.endswith("\r"):
            text = text[:-1]
            self._pending_cr = True
        normalized = text.replace("\r\n", "\n").replace("\r", "\n")
        cursor = 0
        for match in _NEWLINE_RUN_RE.finditer(normalized):
            self._feed_line_segment(normalized[cursor : match.start()])
            self._discard_line_pending()
            self._add_outer_pending(normalized[match.start() : match.end()])
            cursor = match.end()
        self._feed_line_segment(normalized[cursor:])

    def feed(self, text: str) -> None:
        """Consume decoded text using only chunk-sized temporary allocations."""

        for offset in range(0, len(text), _NORMALIZE_CHUNK_CHARS):
            self._feed_chunk(text[offset : offset + _NORMALIZE_CHUNK_CHARS])

    def finish(self) -> _BoundedText:
        # Current line whitespace is removed by per-line rstrip; outer pending
        # whitespace is removed by the final strip.
        self._discard_line_pending()
        self._discard_outer_pending()
        return _BoundedText(text="".join(self._parts), char_count=self._char_count)


def _kind(filename: str, content_type: str | None) -> str | None:
    """Resolve an extractor kind ("pdf" | "docx" | "text") from extension, then MIME."""

    ext = splitext((filename or "").lower())[1]
    if ext in _PDF_EXTS:
        return "pdf"
    if ext in _DOCX_EXTS:
        return "docx"
    if ext in _TEXT_EXTS:
        return "text"
    mime = (content_type or "").split(";", 1)[0].strip().lower()
    if mime in _PDF_MIMES:
        return "pdf"
    if mime in _DOCX_MIMES:
        return "docx"
    if mime in _TEXT_MIMES or (mime.startswith("text/") and mime != "text/"):
        return "text"
    return None


def _complex_format_budget(max_chars: int) -> int:
    return max_chars if max_chars > 0 else _UNBOUNDED_COMPLEX_FORMAT_CHARS


def _extraction_worker_environment() -> dict[str, str]:
    """Return the complete, intentionally minimal parser-worker environment."""

    return dict(_EXTRACTION_WORKER_ENV)


def _configure_pypdf_limits() -> None:
    """Lower pypdf's process-local decompression ceilings before parsing input."""

    from pypdf import filters
    from pypdf.generic import _data_structures

    for name in (
        "FLATE_MAX_BUFFER_SIZE",
        "JBIG2_MAX_OUTPUT_LENGTH",
        "LZW_MAX_OUTPUT_LENGTH",
        "MAX_ARRAY_BASED_STREAM_OUTPUT_LENGTH",
        "MAX_DECLARED_STREAM_LENGTH",
        "RUN_LENGTH_MAX_OUTPUT_LENGTH",
        "ZLIB_MAX_OUTPUT_LENGTH",
    ):
        current = int(getattr(filters, name))
        if current == 0 or current > _MAX_PDF_DECODED_STREAM_BYTES:
            setattr(filters, name, _MAX_PDF_DECODED_STREAM_BYTES)
    _data_structures.CONTENT_STREAM_ARRAY_MAX_LENGTH = min(
        int(_data_structures.CONTENT_STREAM_ARRAY_MAX_LENGTH),
        _MAX_PDF_CONTENT_STREAMS,
    )


def _extract_pdf_in_worker(data: bytes, *, max_chars: int) -> _BoundedText:
    """Parse one PDF inside the resource-limited child process."""

    from pypdf import PasswordType, PdfReader

    _configure_pypdf_limits()
    reader = PdfReader(BytesIO(data))
    if reader.is_encrypted:
        decrypt_result: Any = None
        try:
            decrypt_result = reader.decrypt("")
        except Exception:
            pass
        # pypdf returns PasswordType.NOT_DECRYPTED (integer value 0) for a
        # wrong password; it does not consistently raise. Accept the library's
        # two successful enum values plus exact positive built-ins used by
        # compatible releases/mocks, without coercing an arbitrary result.
        decrypt_succeeded = (
            type(decrypt_result) is PasswordType
            and decrypt_result in {PasswordType.USER_PASSWORD, PasswordType.OWNER_PASSWORD}
        ) or (type(decrypt_result) in {bool, int} and decrypt_result > 0)
        if not decrypt_succeeded:
            raise _ExtractionPolicyError(
                _WORKER_ERROR_PASSWORD_PROTECTED,
                "PDF is password-protected",
            )
    if len(reader.pages) > _MAX_PDF_PAGES:
        raise _ExtractionPolicyError(
            _WORKER_ERROR_SAFETY_LIMIT,
            f"PDF exceeds the {_MAX_PDF_PAGES}-page extraction limit",
        )

    budget = _complex_format_budget(max_chars)
    builder = _NormalizedTextBuilder(budget)
    decoded_content_bytes = 0
    for page_number, page in enumerate(reader.pages):
        contents = page.get_contents()
        if contents is not None:
            decoded_content_bytes += len(contents.get_data())
            if decoded_content_bytes > _MAX_PDF_DECODED_CONTENT_BYTES:
                raise _ExtractionPolicyError(
                    _WORKER_ERROR_SAFETY_LIMIT,
                    "PDF decoded page content exceeds the extraction memory limit",
                )
        if page_number:
            builder.feed("\n\n")
        builder.feed(page.extract_text() or "")
        if builder.char_count > budget:
            break
    return builder.finish()


def _worker_failure_message(payload: bytes, *, label: str) -> str:
    """Map an untrusted worker payload to a stable server-owned diagnostic."""

    try:
        code = payload[1 : 1 + _WORKER_ERROR_BYTES].decode("ascii")
    except UnicodeDecodeError:
        code = ""
    if code == _WORKER_ERROR_SAFETY_LIMIT:
        return f"{label} extraction exceeded a safety limit"
    if code == _WORKER_ERROR_INPUT_LIMIT:
        return f"{label} extraction input exceeds its hard limit"
    if code == _WORKER_ERROR_CONFIGURATION:
        return f"{label} extraction worker configuration failed"
    if code == _WORKER_ERROR_PASSWORD_PROTECTED and label == "PDF":
        return "PDF is password-protected"
    # Unknown codes and parser failures are intentionally indistinguishable.
    # Parser exception text can contain arbitrary attacker-controlled document
    # content and must never cross into ExtractionResult or durable parse_error.
    return f"{label} extraction failed"


def _decode_pdf_worker_response(payload: bytes, *, max_chars: int) -> _BoundedText:
    if not payload:
        raise ValueError("PDF extraction worker exited without a result")
    status = payload[:1]
    if status == b"E":
        raise ValueError(_worker_failure_message(payload, label="PDF"))
    if status != b"S" or len(payload) < 9:
        raise ValueError("PDF extraction worker returned a malformed result")

    char_count = int.from_bytes(payload[1:9], "big", signed=False)
    text_bytes = payload[9:]
    budget = _complex_format_budget(max_chars)
    if len(text_bytes) > budget * 4:
        raise ValueError("PDF extraction worker exceeded its output limit")
    try:
        text: str | None = text_bytes.decode("utf-8")
    except UnicodeDecodeError:
        # UnicodeDecodeError retains the complete worker buffer on ``.object``.
        # Raise only after leaving the handler so neither cause nor context can
        # carry extracted document content across the service boundary.
        text = None
    if text is None:
        raise ValueError("PDF extraction worker returned invalid UTF-8")
    if len(text) > budget or char_count < len(text):
        raise ValueError("PDF extraction worker returned inconsistent bounds")
    return _BoundedText(text=text, char_count=char_count)


def _extract_pdf(data: bytes, *, max_chars: int) -> _BoundedText:
    """Run pypdf with decoded-byte, address-space, CPU, and wall-time bounds."""

    if len(data) > _MAX_PDF_INPUT_BYTES:
        raise ValueError(f"PDF exceeds the {_MAX_PDF_INPUT_BYTES}-byte extraction limit")
    command = [
        sys.executable,
        "-I",
        "-m",
        "apex.services.pdf_extraction_worker",
        str(max_chars),
        str(_PDF_WORKER_MEMORY_BYTES),
        str(_PDF_WORKER_CPU_SECONDS),
        str(_MAX_PDF_INPUT_BYTES),
    ]
    process = subprocess.Popen(  # noqa: S603 - fixed interpreter/module, no shell
        command,
        stdin=subprocess.PIPE,
        stdout=subprocess.PIPE,
        stderr=subprocess.DEVNULL,
        env=_extraction_worker_environment(),
    )
    timed_out = False
    try:
        stdout, _stderr = process.communicate(input=data, timeout=_PDF_WORKER_WALL_SECONDS)
    except subprocess.TimeoutExpired:
        process.kill()
        try:
            process.communicate(timeout=_PDF_WORKER_REAP_SECONDS)
        except subprocess.TimeoutExpired:
            # kill() is definitive on supported deployment platforms; a second
            # call is harmless and communicate() reaps as soon as the kernel does.
            process.kill()
            process.communicate()
        timed_out = True
        stdout = b""
    if timed_out:
        # TimeoutExpired can retain partial worker stdout on ``.output``.  Do
        # not chain it into the stable extraction failure.
        raise ValueError(f"PDF extraction exceeded the {_PDF_WORKER_WALL_SECONDS:g}-second limit")
    if process.returncode != 0:
        raise ValueError("PDF extraction exceeded its CPU or memory limit")
    return _decode_pdf_worker_response(stdout, max_chars=max_chars)


def _extract_docx_in_worker(data: bytes, *, max_chars: int) -> _BoundedText:
    """Parse one DOCX inside the resource-limited child process."""

    from docx import Document as DocxDocument

    with zipfile.ZipFile(BytesIO(data)) as archive:
        members = archive.infolist()
        if len(members) > _MAX_DOCX_ENTRIES:
            raise _ExtractionPolicyError(
                _WORKER_ERROR_SAFETY_LIMIT,
                "DOCX contains too many archive entries",
            )
        expanded_limit = min(
            _MAX_DOCX_EXPANDED_BYTES,
            max(8 * 1024 * 1024, max_chars * 32 if max_chars > 0 else 0),
        )
        expanded = sum(member.file_size for member in members)
        if expanded > expanded_limit:
            raise _ExtractionPolicyError(
                _WORKER_ERROR_SAFETY_LIMIT,
                "DOCX expanded content exceeds the extraction limit",
            )
        for member in members:
            if member.file_size > 0 and member.compress_size == 0:
                raise _ExtractionPolicyError(
                    _WORKER_ERROR_SAFETY_LIMIT,
                    "DOCX contains an invalid compressed entry",
                )
            if (
                member.compress_size > 0
                and member.file_size / member.compress_size > _MAX_DOCX_COMPRESSION_RATIO
            ):
                raise _ExtractionPolicyError(
                    _WORKER_ERROR_SAFETY_LIMIT,
                    "DOCX archive compression ratio exceeds the safety limit",
                )

    document = DocxDocument(BytesIO(data))
    builder = _NormalizedTextBuilder(_complex_format_budget(max_chars))
    first_block = True

    def begin_block() -> None:
        nonlocal first_block
        if not first_block:
            builder.feed("\n")
        first_block = False

    for paragraph in document.paragraphs:
        begin_block()
        builder.feed(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            begin_block()
            for cell_number, cell in enumerate(row.cells):
                if cell_number:
                    builder.feed("\t")
                # Feed each cell directly; never construct a row-sized temporary
                # string before the configured retention cap is applied.
                builder.feed(cell.text)
    return builder.finish()


def _decode_docx_worker_response(payload: bytes, *, max_chars: int) -> _BoundedText:
    if not payload:
        raise ValueError("DOCX extraction worker exited without a result")
    status = payload[:1]
    if status == b"E":
        raise ValueError(_worker_failure_message(payload, label="DOCX"))
    if status != b"S" or len(payload) < 9:
        raise ValueError("DOCX extraction worker returned a malformed result")

    char_count = int.from_bytes(payload[1:9], "big", signed=False)
    text_bytes = payload[9:]
    budget = _complex_format_budget(max_chars)
    if len(text_bytes) > budget * 4:
        raise ValueError("DOCX extraction worker exceeded its output limit")
    try:
        text: str | None = text_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = None
    if text is None:
        raise ValueError("DOCX extraction worker returned invalid UTF-8")
    if len(text) > budget or char_count < len(text):
        raise ValueError("DOCX extraction worker returned inconsistent bounds")
    return _BoundedText(text=text, char_count=char_count)


def _extract_docx(data: bytes, *, max_chars: int) -> _BoundedText:
    """Run python-docx/lxml with address-space, CPU, and wall-time bounds."""

    if len(data) > _MAX_DOCX_INPUT_BYTES:
        raise ValueError(f"DOCX exceeds the {_MAX_DOCX_INPUT_BYTES}-byte extraction limit")
    command = [
        sys.executable,
        "-I",
        "-m",
        "apex.services.docx_extraction_worker",
        str(max_chars),
        str(_DOCX_WORKER_MEMORY_BYTES),
        str(_DOCX_WORKER_CPU_SECONDS),
        str(_MAX_DOCX_INPUT_BYTES),
    ]
    process = subprocess.Popen(  # noqa: S603 - fixed interpreter/module, no shell
        command,
        stdin=subprocess.PIPE,
        stdout=subprocess.PIPE,
        stderr=subprocess.DEVNULL,
        env=_extraction_worker_environment(),
    )
    timed_out = False
    try:
        stdout, _stderr = process.communicate(input=data, timeout=_DOCX_WORKER_WALL_SECONDS)
    except subprocess.TimeoutExpired:
        process.kill()
        try:
            process.communicate(timeout=_DOCX_WORKER_REAP_SECONDS)
        except subprocess.TimeoutExpired:
            process.kill()
            process.communicate()
        timed_out = True
        stdout = b""
    if timed_out:
        raise ValueError(f"DOCX extraction exceeded the {_DOCX_WORKER_WALL_SECONDS:g}-second limit")
    if process.returncode != 0:
        raise ValueError("DOCX extraction exceeded its CPU or memory limit")
    return _decode_docx_worker_response(stdout, max_chars=max_chars)


def _extract_plaintext(data: bytes, *, max_chars: int) -> _BoundedText:
    """Incrementally decode and normalize a bounded prefix of arbitrary UTF-8."""

    builder = _NormalizedTextBuilder(max_chars if max_chars > 0 else None)
    decoder = codecs.getincrementaldecoder("utf-8")(errors="replace")
    for offset in range(0, len(data), _DECODE_CHUNK_BYTES):
        builder.feed(decoder.decode(data[offset : offset + _DECODE_CHUNK_BYTES], final=False))
    builder.feed(decoder.decode(b"", final=True))
    return builder.finish()


def extract_text(
    data: bytes,
    *,
    filename: str,
    content_type: str | None,
    max_chars: int,
) -> ExtractionResult:
    """Best-effort bounded extraction. Content failures are reported, not raised."""

    kind = _kind(filename, content_type)
    if kind is None:
        return ExtractionResult(text="", status=PARSE_UNSUPPORTED, char_count=0)
    try:
        if kind == "pdf":
            extracted = _extract_pdf(data, max_chars=max_chars)
        elif kind == "docx":
            extracted = _extract_docx(data, max_chars=max_chars)
        else:
            extracted = _extract_plaintext(data, max_chars=max_chars)
    except Exception as exc:
        return ExtractionResult(
            text="",
            status=PARSE_FAILED,
            char_count=0,
            error=bounded_diagnostic(exc, max_chars=500),
        )

    text = extracted.text
    full_len = extracted.char_count
    if max_chars > 0 and full_len > max_chars:
        text = text[:max_chars].rstrip() + f"\n\n…[truncated {full_len - max_chars} characters]"
    return ExtractionResult(text=text, status=PARSE_PARSED, char_count=full_len)


def derive_summary(text: str, *, max_chars: int) -> str | None:
    """Return a whitespace-collapsed first non-empty paragraph."""

    for block in text.split("\n\n"):
        candidate = " ".join(block.split())
        if candidate:
            if max_chars > 0 and len(candidate) > max_chars:
                return candidate[: max_chars - 1].rstrip() + "…"
            return candidate
    return None
